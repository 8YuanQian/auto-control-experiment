"""
update_report.py
用 MATLAB/Simulink 仿真数据和截图优化课设报告
策略：在关键段落之后插入截图 + 更新表格数据 + 修复格式字体
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

# === 路径配置 ===
DESKTOP = os.path.join(os.environ['USERPROFILE'], 'Desktop')
LESSON_DIR = os.path.join(DESKTOP, 'professional lesson')
REPORT_PATH = os.path.join(LESSON_DIR, '四旋翼无人机姿态稳定控制系统设计_课设报告_完整版.docx')
SCREENSHOTS_DIR = os.path.join(LESSON_DIR, 'matlab_sim', 'screenshots')
OUTPUT_PATH = os.path.join(DESKTOP, '四旋翼无人机姿态控制_课设报告_Simulink版.docx')

# === 截图映射 ===
SCREENSHOTS = [
    ('01_open_loop_analysis.png',      '开环传递函数为',    '图4-2  滚转通道开环特性分析（根轨迹/Bode图/阶跃响应/零极点分布）', 5.8),
    ('02_pid_design_analysis.png',     '4.4.1 PID',         '图4-3  PID控制器设计分析（补偿后根轨迹/Bode图/闭环阶跃/零极点）', 5.8),
    ('03_lqr_design_analysis.png',     '4.4.2 LQR',         '图4-4  LQR控制器设计分析（闭环阶跃/极点分布/初始响应/反馈增益）', 5.8),
    ('04_pid_simulink_model.png',      'Simulink建模与参数整定', '图5-1  PID控制Simulink仿真模型（Step→PID→Saturation→Plant+Disturbance闭环）', 5.5),
    ('10_code_pid_design.png',         'To Workspace导出phi', '图5-2  PID控制器MATLAB设计代码（根轨迹法整定，含性能指标计算）', 5.5),
    ('05_lqr_simulink_model.png',      'LQR控制器设计与Simulink实现', '图5-3  LQR控制Simulink仿真模型（前馈+State-Space+LQR Gain+状态估计）', 5.5),
    ('11_code_lqr_design.png',         '全状态反馈',        '图5-4  LQR控制器MATLAB设计代码（状态空间建模+lqr()求解+闭环仿真）', 5.5),
    ('12_code_simulink_build.png',     'Simulink Library',  '图5-5  Simulink模型编程构建代码（add_block/set_param/add_line自动化建模）', 5.5),
    ('07_pid_lqr_comparison.png',      '6.1 标称性能',      '图6-1  PID vs LQR阶跃响应全面对比（滚转角/控制量/超调放大/性能指标表）', 5.8),
    ('08_disturbance_comparison.png',  '6.2 抗扰动',        '图6-2  抗扰动性能对比（t=5s阵风脉冲0.05Nm，PID最大偏差4.7°/LQR仅3.5°）', 5.8),
    ('09_cascade_analysis.png',        '6.3.4 双回路',      '图6-3  双回路串级控制架构性能分析（阶跃响应/抗扰动/四项指标柱状对比）', 5.8),
    ('06_cascade_simulink_model.png',  '串级控制架构框图',  '图6-4  双回路串级控制Simulink仿真模型（外环P+内环PD+Derivative速率反馈）', 5.5),
]


def set_run_font(run, cn='宋体', en='Times New Roman', size=Pt(12), bold=None):
    """设置 run 的中英文字体"""
    run.font.name = en
    run.font.size = size
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    # 移除旧字体设置
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:cs'), en)
    rPr.insert(0, rFonts)


def add_image_at_position(doc, image_path, width_inches, after_para_idx):
    """
    在指定段落索引后插入图片。
    做法：先在末尾添加带图片的段落（让 python-docx 处理关系），
    再把生成的 XML 元素搬移到目标位置，最后删除末尾的临时段落。
    """
    # 获取图片宽高比
    with Image.open(image_path) as img:
        img_w, img_h = img.size
    aspect = img_h / img_w

    # 1. 在文档末尾创建临时段落来承载图片
    tmp_para = doc.add_paragraph()
    tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tmp_run = tmp_para.add_run()
    tmp_run.add_picture(image_path, width=Inches(width_inches))

    # 获取刚创建的段落 element
    img_elem = tmp_para._element

    # 2. 创建题注段落
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run()
    # 取对应的 caption
    for fn, sk, caption, w in SCREENSHOTS:
        if fn in image_path:
            cap_run.text = caption
            break
    set_run_font(cap_run, '宋体', 'Times New Roman', Pt(10.5))
    cap_elem = cap_para._element

    # 创建末尾空行
    spacer_para = doc.add_paragraph()
    spacer_para.add_run().text = ''
    spacer_elem = spacer_para._element

    # 3. 找到目标段落 element
    target_elem = doc.paragraphs[after_para_idx]._element
    parent = target_elem.getparent()

    # 找到目标在 parent 中的位置
    target_pos = list(parent).index(target_elem)

    # 4. 从末尾移除这些元素，插入到目标之后
    parent.remove(img_elem)
    parent.remove(cap_elem)
    parent.remove(spacer_elem)

    # 倒序插入保序（图片、题注、空行）
    parent.insert(target_pos + 1, spacer_elem)
    parent.insert(target_pos + 1, cap_elem)
    parent.insert(target_pos + 1, img_elem)

    return True


# === 主流程 ===
print('=' * 60)
print('四旋翼无人机课设报告优化脚本')
print('=' * 60)

print('\n>>> 1. 读取原始报告...')
doc = Document(REPORT_PATH)
print(f'  段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)}')

# === 2. 插入截图 ===
print('\n>>> 2. 插入 Simulink 仿真截图...')

# 先记录原始段落数量，方便后续索引
# 因为插入操作会改变段落顺序，我们需要从后往前插入
# 找到所有插入点
insertion_plan = []  # (image_path, width, target_index, filename)
for filename, search_key, caption, width in SCREENSHOTS:
    img_path = os.path.join(SCREENSHOTS_DIR, filename)
    if not os.path.exists(img_path):
        print(f'  [SKIP] 文件不存在: {filename}')
        continue

    found = False
    for i, para in enumerate(doc.paragraphs):
        if search_key in para.text:
            insertion_plan.append((img_path, width, i, filename, caption))
            found = True
            break

    if not found:
        print(f'  [WARN] 未找到插入位置: "{search_key[:40]}"')

# 按段落索引从大到小排序（从后往前插入，避免索引偏移）
insertion_plan.sort(key=lambda x: x[2], reverse=True)

inserted = 0
for img_path, width, target_idx, filename, caption in insertion_plan:
    try:
        # 临时替换 caption 查找 - 用 filename 匹配
        add_image_at_position(doc, img_path, width, target_idx)
        print(f'  [OK] {filename} → P{target_idx}')
        inserted += 1
    except Exception as e:
        print(f'  [ERR] {filename}: {e}')

print(f'  成功插入 {inserted}/{len(SCREENSHOTS)} 张截图')

# === 3. 修复全文正文字体 ===
print('\n>>> 3. 修复全文字体格式...')

CN_BODY = '宋体'
EN_BODY = 'Times New Roman'
CN_HEADING = '黑体'

HEADING_KW = ['背景及意义', '国内外研究现状', '问题描述与建模', '总体方案',
              '姿态动力学数学模型', '控制器设计', '性能指标',
              '基于Simulink', '仿真结果与分析', '总结',
              '四旋翼物理参数', '仿真环境参数', '创新改进',
              '前馈补偿', '增益调度', '抗积分饱和', '双回路串级',
              '摘要', '关键词', 'Abstract', 'Key words', '参考文献',
              '标称性能仿真', '抗扰动仿真']

body_fixed = 0
heading_fixed = 0

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    is_heading = False
    for hw in HEADING_KW:
        if hw in text and len(text) < 80:
            is_heading = True
            break
    # 数字编号开头 + 较短
    if not is_heading and len(text) < 50:
        for prefix in ['1.', '2.', '3.', '4.', '5.', '6.', '7.',
                       '4.1', '4.2', '4.3', '4.4', '4.5',
                       '5.1', '5.2', '5.3',
                       '6.1', '6.2', '6.3']:
            if text.startswith(prefix):
                is_heading = True
                break

    cn_font = CN_HEADING if is_heading else CN_BODY
    size = Pt(14) if is_heading else Pt(12)

    for run in para.runs:
        if run.text.strip():
            rPr = run._element.find(qn('w:rPr'))
            current_cn = None
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    current_cn = rFonts.get(qn('w:eastAsia'))

            if current_cn is None or current_cn != cn_font:
                set_run_font(run, cn_font, EN_BODY, size, run.bold)
                if is_heading:
                    heading_fixed += 1
                else:
                    body_fixed += 1

print(f'  标题字体修复: {heading_fixed} runs')
print(f'  正文字体修复: {body_fixed} runs')

# === 4. 修复表格字体 ===
print('\n>>> 4. 修复表格字体...')
for ti, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        set_run_font(run, '宋体', 'Times New Roman', Pt(10.5))
    print(f'  表格{ti+1}: {len(table.rows)}行x{len(table.columns)}列 已修复')

# === 5. 页面信息 ===
print('\n>>> 5. 页面设置...')
for si, section in enumerate(doc.sections):
    print(f'  节{si+1}: {section.page_width.inches:.1f}x{section.page_height.inches:.1f} in, '
          f'边距 L={section.left_margin.inches:.2f} R={section.right_margin.inches:.2f}')

# === 6. 保存 ===
print(f'\n>>> 6. 保存...')
doc.save(OUTPUT_PATH)

size_kb = os.path.getsize(OUTPUT_PATH) / 1024
print(f'  文件大小: {size_kb:.1f} KB')
print('\n' + '=' * 60)
print('报告优化完成！')
print('输出:', OUTPUT_PATH)
print('=' * 60)
